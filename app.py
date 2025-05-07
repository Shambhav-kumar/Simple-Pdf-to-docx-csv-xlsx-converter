from flask import Flask, request, jsonify, send_from_directory, render_template, send_file
import os
import uuid
import io
from werkzeug.utils import secure_filename
from pdf2docx import Converter
import camelot
import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import tempfile
import logging

app = Flask(__name__)

# Configuration
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['CONVERTED_FOLDER'] = 'converted'
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB limit

# Set up logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Ensure directories exist
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
os.makedirs(app.config['CONVERTED_FOLDER'], exist_ok=True)

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() == 'pdf'

def convert_academic_pdf_to_docx(pdf_path: str, docx_path: str) -> bool:
    """Enhanced PDF to DOCX conversion for academic documents"""
    try:
        # Use temporary file for safer conversion
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
            temp_docx_path = temp_file.name
        
        # Convert with pdf2docx with academic document optimizations
        cv = Converter(pdf_path)
        
        # Custom conversion parameters for academic documents
        cv.convert(
            temp_docx_path,
            start=0,
            end=None,
            multi_processing=True,
            debug=False,
            ignore_bad_chars=True,  # Handle special characters better
            keep_blank_chars=True,  # Preserve spaces
            layout_analysis=True,   # Better structure recognition
            table_detection=True,   # Improved table handling
            complex_columns=True    # Better multi-column handling
        )
        cv.close()
        
        # Post-processing to improve academic document formatting
        doc = Document(temp_docx_path)
        
        # Set default font for the document
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(12)
        
        # Save the improved document
        doc.save(docx_path)
        
        # Clean up temporary file
        if os.path.exists(temp_docx_path):
            os.unlink(temp_docx_path)
            
        return True
        
    except Exception as e:
        logger.error(f"Academic PDF conversion error: {str(e)}")
        if 'temp_docx_path' in locals() and os.path.exists(temp_docx_path):
            os.unlink(temp_docx_path)
        return False

@app.route('/')
def home():
    return render_template('index.html')

@app.route('/convert', methods=['POST'])
def convert_file():
    if 'file' not in request.files:
        return jsonify({'error': 'No file uploaded'}), 400

    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'No file selected'}), 400

    if not allowed_file(file.filename):
        return jsonify({'error': 'Only PDF files allowed'}), 400

    format_type = request.form.get('format')
    if format_type not in ['docx', 'csv', 'xlsx']:
        return jsonify({'error': 'Invalid format'}), 400

    try:
        # Save uploaded file
        filename = secure_filename(file.filename)
        unique_id = uuid.uuid4().hex
        upload_path = os.path.join(app.config['UPLOAD_FOLDER'], f"{unique_id}_{filename}")
        file.save(upload_path)

        # Generate output filename
        output_filename = f"{os.path.splitext(filename)[0]}.{format_type}"
        output_path = os.path.join(app.config['CONVERTED_FOLDER'], f"{unique_id}_{output_filename}")

        # Conversion
        if format_type == 'docx':
            if not convert_academic_pdf_to_docx(upload_path, output_path):
                return jsonify({'error': 'Failed to convert academic PDF to DOCX'}), 500
        else:
            tables = camelot.read_pdf(upload_path, pages='all')
            if tables.n == 0:
                return jsonify({'error': 'No tables found in PDF'}), 400
            df = pd.concat([table.df for table in tables]) if tables.n > 1 else tables[0].df
            if format_type == 'csv':
                df.to_csv(output_path, index=False)
            else:
                df.to_excel(output_path, index=False)

        return jsonify({
            'filename': output_filename,
            'download_url': f"/download/{unique_id}_{output_filename}",
            'preview_url': f"/preview_output/{unique_id}_{output_filename}",
            'format': format_type
        })

    except Exception as e:
        logger.error(f"Conversion error: {str(e)}")
        return jsonify({'error': str(e)}), 500
    finally:
        # Clean up uploaded file
        if 'upload_path' in locals() and os.path.exists(upload_path):
            try:
                os.unlink(upload_path)
            except Exception as e:
                logger.error(f"Error cleaning up file: {str(e)}")

@app.route('/download/<filename>')
def download_file(filename):
    return send_from_directory(app.config['CONVERTED_FOLDER'], filename, as_attachment=True)

@app.route('/preview_output/<filename>')
def preview_output_file(filename):
    filepath = os.path.join(app.config['CONVERTED_FOLDER'], filename)
    ext = filename.rsplit('.', 1)[1].lower()
    
    if ext == 'docx':
        try:
            doc = Document(filepath)
            html_content = """
            <!DOCTYPE html>
            <html>
            <head>
                <title>DOCX Preview</title>
                <style>
                    body { font-family: 'Times New Roman', serif; line-height: 1.6; padding: 20px; }
                    .docx-wrapper { max-width: 800px; margin: 0 auto; }
                    table { border-collapse: collapse; width: 100%; margin: 20px 0; }
                    table, th, td { border: 1px solid #ddd; }
                    th, td { padding: 8px; text-align: left; }
                    h1, h2, h3 { color: #2c3e50; }
                    h1 { font-size: 18pt; }
                    h2 { font-size: 16pt; }
                    p { margin: 10px 0; font-size: 12pt; }
                </style>
            </head>
            <body>
                <div class="docx-wrapper">
            """
            
            for para in doc.paragraphs:
                html_content += f"<p>{para.text}</p>"
            
            for table in doc.tables:
                html_content += "<table>"
                for row in table.rows:
                    html_content += "<tr>"
                    for cell in row.cells:
                        html_content += f"<td>{cell.text}</td>"
                    html_content += "</tr>"
                html_content += "</table>"
            
            html_content += """
                </div>
            </body>
            </html>
            """
            return html_content
        except Exception as e:
            logger.error(f"Preview error: {str(e)}")
            return f"Error generating preview: {str(e)}", 500
    elif ext == 'csv':
        # Read CSV and convert to HTML table
        df = pd.read_csv(filepath)
        return df.to_html(classes='data-table', index=False)
    elif ext == 'xlsx':
        # Read Excel and convert to HTML table
        df = pd.read_excel(filepath)
        return df.to_html(classes='data-table', index=False)
    else:
        return "Preview not available for this file type"

if __name__ == '__main__':
    app.run(debug=True)